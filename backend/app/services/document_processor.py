"""
EZ-NEXUS AI — Document Processor Service
Extracts text from PDF, Word, Excel/CSV, and image files.
All external libraries (pypdf, pytesseract, etc.) are imported lazily
so the server starts even if OCR dependencies aren't installed.
"""

from __future__ import annotations
import logging
from pathlib import Path
from typing import Tuple

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Unified text extractor for all supported document types.
    Returns (extracted_text, confidence_score 0.0–1.0).
    Confidence reflects how reliable the extraction method is.
    """

    def extract_text(self, path: str, content_type: str | None = None) -> Tuple[str, float]:
        """Main entry point. Routes to the appropriate extractor."""
        file_path = Path(path)
        suffix = file_path.suffix.lower()

        extractors = {
            ".pdf":  self._extract_pdf,
            ".docx": self._extract_docx,
            ".doc":  self._extract_docx,
            ".xlsx": self._extract_spreadsheet,
            ".xls":  self._extract_spreadsheet,
            ".csv":  self._extract_spreadsheet,
            ".png":  self._extract_image_ocr,
            ".jpg":  self._extract_image_ocr,
            ".jpeg": self._extract_image_ocr,
            ".tif":  self._extract_image_ocr,
            ".tiff": self._extract_image_ocr,
            ".heic": self._extract_image_ocr,
        }

        extractor = extractors.get(suffix)
        if extractor is None:
            logger.warning("Unsupported file type: %s", suffix)
            return "", 0.0

        try:
            return extractor(file_path)
        except Exception as e:
            logger.error("Extraction failed for %s: %s", path, e)
            return "", 0.0

    # ── PDF ───────────────────────────────────────────────────────────────────

    def _extract_pdf(self, path: Path) -> Tuple[str, float]:
        try:
            from pypdf import PdfReader
        except ImportError:
            logger.warning("pypdf not installed — PDF extraction unavailable")
            return f"[PDF extraction requires pypdf: pip install pypdf]\nFilename: {path.name}", 0.1

        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            try:
                text = page.extract_text() or ""
                pages.append(text)
            except Exception:
                pass

        text = "\n".join(pages).strip()
        if text and len(text) > 50:
            return text, 0.88

        # PDF appears to be a scan — try OCR
        logger.info("PDF has no embedded text — attempting OCR: %s", path.name)
        return self._ocr_pdf_pages(path)

    def _ocr_pdf_pages(self, path: Path) -> Tuple[str, float]:
        """Convert PDF pages to images and run Tesseract OCR."""
        try:
            from pdf2image import convert_from_path
            import pytesseract
            images = convert_from_path(str(path), dpi=200)
            pages = [pytesseract.image_to_string(img) for img in images]
            text = "\n".join(pages).strip()
            return text, 0.65 if text else 0.1
        except ImportError:
            return "[Scanned PDF detected but pdf2image/pytesseract not installed]", 0.1
        except Exception as e:
            logger.error("PDF OCR failed: %s", e)
            return "", 0.1

    # ── Word ──────────────────────────────────────────────────────────────────

    def _extract_docx(self, path: Path) -> Tuple[str, float]:
        try:
            from docx import Document
        except ImportError:
            return "[Word extraction requires python-docx: pip install python-docx]", 0.1

        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())

        return "\n".join(parts), 0.92

    # ── Spreadsheet / CSV ─────────────────────────────────────────────────────

    def _extract_spreadsheet(self, path: Path) -> Tuple[str, float]:
        try:
            import pandas as pd
        except ImportError:
            return "[Spreadsheet extraction requires pandas: pip install pandas openpyxl]", 0.1

        try:
            if path.suffix.lower() == ".csv":
                df = pd.read_csv(path, encoding="utf-8", encoding_errors="replace")
            else:
                df = pd.read_excel(path)

            # Return CSV string representation — preserves structure
            text = df.to_csv(index=False)
            return text, 0.95
        except Exception as e:
            logger.error("Spreadsheet extraction error: %s", e)
            return "", 0.1

    # ── Image OCR ─────────────────────────────────────────────────────────────

    def _extract_image_ocr(self, path: Path) -> Tuple[str, float]:
        try:
            from PIL import Image
            import pytesseract
        except ImportError:
            return "[Image OCR requires pillow + pytesseract: pip install pillow pytesseract]", 0.1

        try:
            image = Image.open(str(path))
            # Convert HEIC or other formats to RGB
            if image.mode not in ("RGB", "L"):
                image = image.convert("RGB")

            # Try to improve OCR quality: resize if too small
            w, h = image.size
            if w < 1200:
                scale = 1200 / w
                image = image.resize((int(w * scale), int(h * scale)))

            text = pytesseract.image_to_string(image, config="--psm 6")
            confidence = 0.72 if len(text.strip()) > 50 else 0.35
            return text.strip(), confidence
        except Exception as e:
            logger.error("Image OCR error: %s", e)
            return "", 0.1

    # ── Utilities ─────────────────────────────────────────────────────────────

    @staticmethod
    def classify_document_type(filename: str, text: str = "") -> str:
        """Guess document type from filename and content keywords."""
        name = filename.lower()
        content = text.lower()

        if any(x in name for x in [".pdf"]):
            if any(x in content for x in ["insurance", "member id", "payer", "plan"]):
                return "insurance_card"
            if any(x in content for x in ["patient", "date of birth", "dob", "diagnosis"]):
                return "patient_form"
            if any(x in content for x in ["invoice", "receipt", "total", "subtotal"]):
                return "invoice_receipt"
            return "pdf_document"

        if any(x in name for x in [".docx", ".doc"]):
            return "word_document"

        if any(x in name for x in [".xlsx", ".xls", ".csv"]):
            return "spreadsheet"

        if any(x in name for x in [".png", ".jpg", ".jpeg", ".heic", ".tif"]):
            if any(x in content for x in ["insurance", "member", "rx"]):
                return "insurance_card_scan"
            return "image_scan"

        return "unknown"
